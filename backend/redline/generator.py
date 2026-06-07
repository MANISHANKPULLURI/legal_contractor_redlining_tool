from docx import Document
from docx.shared import RGBColor
from datetime import datetime





def add_deleted_text(document, text):

    para = document.add_paragraph()

    run = para.add_run(text)

    run.font.strike = True

    run.font.color.rgb = RGBColor(
        180,
        0,
        0
    )







def add_added_text(document, text):

    para = document.add_paragraph()

    run = para.add_run(text)

    run.bold = True

    run.font.color.rgb = RGBColor(
        0,
        120,
        0
    )









def has_real_rewrite(
    original,
    rewritten
):


    if not rewritten:

        return False



    invalid_outputs = [

        "no rewrite required",
        "no rewrite required.",
        "no rewrite required for this clause.",
        "no changes needed",
        "no modification required",
        "no rewritten clause needed"

    ]



    cleaned = rewritten.strip().lower()



    if cleaned in invalid_outputs:

        return False



    if cleaned == original.strip().lower():

        return False



    return True










def create_redline_doc(
    rewrites,
    output_path="redlined_contract.docx"
):


    document = Document()



    document.add_heading(
        "Lexo AI - Contract Review Report",
        level=1
    )



    document.add_paragraph(
        "AI generated contract risk analysis, recommendations, and redlined improvements."
    )



    document.add_paragraph(
        f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    )







    for item in rewrites:


        analysis = item.get(
            "analysis",
            {}
        )




        clause_number = item.get(
            "clause_number",
            "Unknown"
        )




        original_clause = (

            item.get("original_clause")

            or

            item.get("clause")

            or

            ""

        )




        rewritten_clause = (

            item.get("rewritten_clause")

            or

            analysis.get("rewritten_clause")

            or

            ""

        )




        risk_level = (

            item.get("risk_level")

            or

            analysis.get("risk_level")

            or

            "LOW"

        )




        explanation = (

            item.get("explanation")

            or

            analysis.get("explanation")

            or

            ""

        )




        recommendation = (

            item.get("recommendation")

            or

            analysis.get("recommendation")

            or

            "No modification required based on current review criteria."

        )




        issues = (

            item.get("issues")

            or

            analysis.get("issues")

            or

            []

        )







        document.add_page_break()



        document.add_heading(
            f"Clause {clause_number} - AI Legal Review",
            level=2
        )







        # Risk section


        document.add_heading(
            "Risk Analysis",
            level=3
        )



        document.add_paragraph(
            f"Risk Level: {risk_level}"
        )






        if issues:


            document.add_paragraph(
                "Issues Found:"
            )



            for issue in issues:


                if isinstance(issue, dict):

                    issue_text = issue.get(
                        "issue",
                        ""
                    )

                    reason = issue.get(
                        "why_risky",
                        ""
                    )


                else:

                    issue_text = str(issue)

                    reason = ""




                document.add_paragraph(
                    f"- {issue_text}"
                )



                if reason:

                    document.add_paragraph(
                        f"Reason: {reason}"
                    )









        if explanation:


            document.add_paragraph(
                "Explanation:"
            )


            document.add_paragraph(
                explanation
            )






        if recommendation:


            document.add_paragraph(
                "Recommendation:"
            )


            document.add_paragraph(
                recommendation
            )









        # Original


        document.add_heading(
            "Original Clause",
            level=3
        )


        document.add_paragraph(
            original_clause
        )









        # Redline


        document.add_heading(
            "Redline Changes",
            level=3
        )





        if has_real_rewrite(
            original_clause,
            rewritten_clause
        ):


            document.add_paragraph(
                "Removed / Modified Text:"
            )


            add_deleted_text(
                document,
                original_clause
            )



            document.add_paragraph(
                "Suggested Revision:"
            )


            add_added_text(
                document,
                rewritten_clause
            )



            final_clause = rewritten_clause





        else:


            document.add_paragraph(
                "No modification recommended. Clause retained as originally drafted."
            )


            final_clause = original_clause







        document.add_heading(
            "Final Recommended Clause",
            level=3
        )


        document.add_paragraph(
            final_clause
        )








    document.save(
        output_path
    )



    return output_path